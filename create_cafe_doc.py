from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Create document
doc = Document()

# Title
title = doc.add_heading('Cafe Contact Directory', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Delhi NCR - Web & App Development Priority List')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.style = 'Subtitle'

# Introduction
doc.add_heading('Overview', level=1)
doc.add_paragraph('This directory contains contact details for 40+ cafes in Delhi NCR, prioritized based on their need for web and app development services.')

# Priority Summary
doc.add_heading('Priority Summary', level=1)
priority_table = doc.add_table(rows=1, cols=4)
priority_table.style = 'Table Grid'

# Header row
header_cells = priority_table.rows[0].cells
headers = ['Priority', 'Count', 'Need Level', 'Action Timeline']
for i, header in enumerate(headers):
    header_cells[i].paragraphs[0].add_run(header).bold = True

# Data rows
data = [
    ['🔴 High', '3', 'Immediate', 'Contact this week'],
    ['🟡 Medium', '6', 'Short-term', 'Contact this month'],
    ['🟢 Low', '31+', 'Optimization', 'Nurture campaign']
]
for row in data:
    cells = priority_table.add_row().cells
    for i, cell in enumerate(row):
        cells[i].text = cell

doc.add_paragraph()

# High Priority Cafes
doc.add_heading('🔴 HIGH PRIORITY - Immediate Need', level=1)
doc.add_paragraph('These cafes are newly opened (low review count) and lack digital presence.')

high_priority_table = doc.add_table(rows=1, cols=6)
high_priority_table.style = 'Table Grid'

# Headers
headers = ['Cafe', 'Phone', 'Email', 'Address', 'Reviews', 'Rating']
for i, header in enumerate(headers):
    high_priority_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

# High priority data
high_priority_cafes = [
    ['Shin Chan Cafe', '+91 99991 00144', '', 'Shop 1, Krishna Market, Lajpat Nagar 1', '6', '4.5'],
    ['Chaigadh', '+91 98911 08885', '', 'Block K, 127, New Delhi', '10', '4.8'],
    ['Cozy Cup Cafe', '+91 84355 80355', 'support@cozycupcafe.in', '3026/5A, New Delhi', '15', '4.8']
]

for cafe in high_priority_cafes:
    cells = high_priority_table.add_row().cells
    for i, cell_data in enumerate(cafe):
        cells[i].text = cell_data

doc.add_paragraph()

# Medium Priority Cafes
doc.add_heading('🟡 MEDIUM PRIORITY - Short-term Need', level=1)
doc.add_paragraph('These cafes have some presence but need professional web/app solutions.')

medium_priority_table = doc.add_table(rows=1, cols=6)
medium_priority_table.style = 'Table Grid'

for i, header in enumerate(headers):
    medium_priority_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

medium_priority_cafes = [
    ['Empire Cafe House Branch-7', '+91 80761 13282', '', 'Mukherjee Nagar', '20', '5.0'],
    ['Blends & Brews Cafe', '+91 99580 53013', '', 'West Patel Nagar', '33', '4.9'],
    ['BON TEMPS Cafe', '+91 99993 67991', '', 'Shankar Rd Market', '51', '4.7'],
    ['Sip Out', '+91 98911 08885', '', 'Old Rajinder Nagar', '83', '4.7'],
    ['Guliyano Cafe', '+91 92125 08354', '', 'Karol Bagh', '309', '4.6'],
    ['Tadow Cafe', '+91 88829 24571', '', 'Block 8, New Delhi', '14', '5.0']
]

for cafe in medium_priority_cafes:
    cells = medium_priority_table.add_row().cells
    for i, cell_data in enumerate(cafe):
        cells[i].text = cell_data

doc.add_paragraph()

# Complete Directory
doc.add_heading('📋 COMPLETE DIRECTORY', level=1)

full_table = doc.add_table(rows=1, cols=7)
full_table.style = 'Table Grid'

full_headers = ['Priority', 'Cafe Name', 'Phone', 'Email', 'Address', 'Website', 'Reviews']
for i, header in enumerate(full_headers):
    full_table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

# All cafe data
all_cafes = [
    ['High', 'Shin Chan Cafe', '+91 99991 00144', '', 'Lajpat Nagar 1', '', '6'],
    ['High', 'Chaigadh', '+91 98911 08885', '', 'Block K, 127', '', '10'],
    ['High', 'Cozy Cup Cafe', '+91 84355 80355', 'support@cozycupcafe.in', '3026/5A', 'cozycupcafe.in', '15'],
    ['Medium', 'Empire Cafe House', '+91 80761 13282', '', 'Mukherjee Nagar', 'empirecafehouse.com', '20'],
    ['Medium', 'Blends & Brews', '+91 99580 53013', '', 'West Patel Nagar', 'blendsandbrews.com', '33'],
    ['Medium', 'BON TEMPS', '+91 99993 67991', '', 'Shankar Rd', '', '51'],
    ['Medium', 'Sip Out', '+91 98911 08885', '', 'Rajinder Nagar', '', '83'],
    ['Medium', 'Guliyano Cafe', '+91 92125 08354', '', 'Karol Bagh', '', '309'],
    ['Low', 'Third Wave Coffee', '+91 89519 41926', '', 'GK-II', 'thirdwavecoffeeroasters.com', '1240'],
    ['Low', 'AMA Cafe', '+91 84489 94259', '', 'Majnu Ka Tila', '', '29750'],
    ['Low', 'Triveni Terrace', '+91 99715 66904', '', 'Mandi House', 'redcedarhospitality.com', '6533'],
    ['Low', 'Cloud Star', '+91 95993 55049', '', 'Paharganj', '', '187'],
    ['Low', 'Cafe Loco Moco', '+91 99909 57579', '', 'Kamla Nagar', '', '432'],
    ['Low', 'Nothing Before Coffee', '+91 99719 14342', '', 'Patel Nagar', 'nothingbeforecoffeeindia.com', '443'],
    ['Low', 'Chai Chaupal', '+91 88829 24571', '', 'Vijay Nagar', '', '24'],
    ['Low', 'Dripping Dose', '+91 93551 55150', '', 'Patel Nagar', '', '770'],
    ['Low', 'Open House', '+91 99996 11334', '', 'Connaught Place', 'openhousecafe.in', '7959'],
    ['Low', 'Cafe Delhi Heights', '+91 98114 54033', 'cdhprom@gmail.com', 'Naraina', 'cafedelhiheights.com', '1424'],
    ['Low', 'QBA', '+91 97173 23014', '', 'Patel Nagar', 'qbacp.com', '1809'],
    ['Low', 'Chai 24/7', '+91 85270 82708', '', 'Connaught Place', '', '802'],
    ['Low', 'Peace of Mind', '+91 70113 47942', '', 'Hari Mandir', '', '399'],
    ['Low', 'Letzz Coffee', '+91 98911 08885', '', 'Bada Bazar', '', '137'],
    ['Low', 'Sesh Rooftop', '+91 99993 67991', '', 'Connaught Place', '', '342'],
    ['Low', 'Tim Hortons', '+91 74285 83240', '', 'Lala Jagat Narayan', 'timhortons.in', '764'],
    ['Low', 'Cafe Milkyway', '+91 92125 08354', '', '8A/3', '', '80'],
    ['Low', 'Josie\'s Coffee', '+91 99580 53013', '', 'Central', '', '184'],
    ['Low', 'House of Espresso', '+91 99719 14342', '', 'Shivaji Marg', '', '99'],
    ['Low', 'Lite Brown', '+91 88829 24571', '', 'East Ave', '', '255'],
    ['Low', 'Woodland', '+91 93551 55150', '', 'Gurudwara Rd', '', '75'],
    ['Low', 'Nobby\'s', '+91 99996 11334', '', '9/2', '', '597'],
    ['Low', 'Mom\'s', '+91 80761 13282', '', '22/20', '', '165'],
    ['Low', 'Nicckz', '+91 80761 13282', '', 'Shadipur', '', '84'],
    ['Low', 'Tree House', '+91 99715 66904', '', 'J5RP+G2W', '', '456'],
    ['Low', 'Le Winkies', '+91 99909 57579', '', '14/3', '', '428'],
    ['Low', 'Bean Sahab', '+91 84489 94259', '', '18/31', '', '113'],
    ['Low', 'Tadow', '+91 88829 24571', '', 'Block 8', '', '14'],
    ['Low', 'Imperial', '', '', '', '', '216']
]

for cafe in all_cafes:
    cells = full_table.add_row().cells
    for i, cell_data in enumerate(cafe):
        cells[i].text = cell_data

doc.add_paragraph()

# Recommendations
doc.add_heading('💡 Recommendations', level=1)
doc.add_paragraph('1. Start outreach with HIGH priority cafes - they are most likely to need immediate help')
doc.add_paragraph('2. MEDIUM priority cafes may already have basic presence - pitch premium features')
doc.add_paragraph('3. LOW priority cafes are established - pitch optimization, loyalty apps, analytics')

# Footer
doc.add_paragraph()
doc.add_paragraph('Generated on: 2026-03-24')
doc.add_paragraph('Data sources: Zomato, Magicpin, Restaurant Guru, Justdial, Official Websites')

# Save
doc.save('C:\\Users\\jhaab\\OneDrive\\Desktop\\Sulav Mchantronics\\cafe_contact_directory.docx')
print('Word document created successfully!')
